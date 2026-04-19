"""
Extractor — trích xuất nội dung văn bản từ nhiều định dạng file.
Hỗ trợ: PDF, TXT, MD, DOCX, CSV, HTML
"""

from __future__ import annotations

import csv
import io
import logging
from pathlib import Path

logger = logging.getLogger(__name__)


class ExtractorError(Exception):
    pass


def extract_text(file_path: str | Path) -> str:
    """
    Nhận đường dẫn file, trả về văn bản thô.
    Tự nhận dạng định dạng qua extension.
    """
    path = Path(file_path)
    ext = path.suffix.lower().lstrip(".")

    extractors = {
        "pdf": _extract_pdf,
        "txt": _extract_text,
        "md": _extract_text,
        "markdown": _extract_text,
        "docx": _extract_docx,
        "doc": _extract_docx,
        "csv": _extract_csv,
        "html": _extract_html,
        "htm": _extract_html,
    }

    extractor = extractors.get(ext)
    if not extractor:
        raise ExtractorError(f"Định dạng '{ext}' chưa được hỗ trợ")

    try:
        return extractor(path)
    except ExtractorError:
        raise
    except Exception as e:
        raise ExtractorError(f"Lỗi khi đọc file {path.name}: {e}") from e


# ── PDF ───────────────────────────────────────────────────────────────────────

def _extract_pdf(path: Path) -> str:
    try:
        import pypdf
    except ImportError:
        raise ExtractorError("Cần cài pypdf: pip install pypdf")

    reader = pypdf.PdfReader(str(path))
    pages: list[str] = []
    for idx, page in enumerate(reader.pages, start=1):
        text = page.extract_text() or ""
        clean_page = text.strip()
        if clean_page:
            pages.append(f"[[PAGE:{idx}]]\n{clean_page}")
    return "\n\n".join(pages)


# ── DOCX ──────────────────────────────────────────────────────────────────────

def _extract_docx(path: Path) -> str:
    try:
        from docx import Document
    except ImportError:
        raise ExtractorError("Cần cài python-docx: pip install python-docx")

    doc = Document(str(path))
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    # Lấy text trong bảng
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                paragraphs.append(" | ".join(cells))
    return "\n\n".join(paragraphs)


# ── TXT / MD ──────────────────────────────────────────────────────────────────

def _extract_text(path: Path) -> str:
    for encoding in ("utf-8", "utf-8-sig", "latin-1"):
        try:
            return path.read_text(encoding=encoding)
        except UnicodeDecodeError:
            continue
    raise ExtractorError(f"Không thể đọc file {path.name} với các encoding phổ biến")


# ── CSV ───────────────────────────────────────────────────────────────────────

def _extract_csv(path: Path) -> str:
    rows: list[str] = []
    with open(path, newline="", encoding="utf-8-sig") as f:
        reader = csv.DictReader(f)
        if reader.fieldnames:
            rows.append(" | ".join(str(h) for h in reader.fieldnames))
            rows.append("-" * 40)
        for row in reader:
            rows.append(" | ".join(str(v) for v in row.values()))
    return "\n".join(rows)


# ── HTML ──────────────────────────────────────────────────────────────────────

def _extract_html(path: Path) -> str:
    try:
        from bs4 import BeautifulSoup
    except ImportError:
        raise ExtractorError("Cần cài beautifulsoup4: pip install beautifulsoup4")

    content = path.read_text(encoding="utf-8", errors="ignore")
    soup = BeautifulSoup(content, "html.parser")
    # Loại bỏ script, style
    for tag in soup(["script", "style", "nav", "footer", "header"]):
        tag.decompose()
    return soup.get_text(separator="\n", strip=True)
